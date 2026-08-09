# generate_documentation.py
# Script untuk membuat dokumen DOCX komprehensif dari monorepo WMS Python

import sys
sys.path.insert(0, r'c:\Users\DELL\Downloads\python\Vibes Coding\wms-monorepo\frontend\api')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── STYLE SETUP ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_code_table(doc, rows_data):
    """Add a two-column table: Code | Explanation"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Baris Kode Python'
    hdr[1].text = 'Penjelasan Bahasa Awam'
    for code, explanation in rows_data:
        row = table.add_row().cells
        row[0].text = code
        row[1].text = explanation
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(10)

# ═══════════════════════════════════════════
# HALAMAN JUDUL
# ═══════════════════════════════════════════
doc.add_paragraph('')
title = doc.add_heading('DOKUMENTASI KOMPREHENSIF', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_heading('WMS Monorepo - Demand Supply Planning System', level=1)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Modul Pembelajaran Mandiri untuk Pemula', style='Intense Quote')
doc.add_paragraph(f'Dihasilkan secara otomatis oleh AI Code Analyst')
doc.add_paragraph('')

# ═══════════════════════════════════════════
# DAFTAR ISI
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('DAFTAR ISI', level=1)
toc_items = [
    'Bab 1: Gambaran Besar Sistem (Bahasa Awam)',
    'Bab 2: Arsitektur & Struktur Kode',
    'Bab 3: Code Review — Temuan & Rekomendasi',
    'Bab 4: Bug Hunting & Edge Cases',
    'Bab 5: Analisis Akurasi Perhitungan',
    'Bab 6: Bedah Kode Per Modul (Line-by-Line)',
    '  6.1 main.py — Pintu Masuk Aplikasi',
    '  6.2 occupancy_engine.py — Mesin MRP & Kapasitas Gudang',
    '  6.3 forecast_engine.py — Mesin Peramalan 16 Model',
    '  6.4 safety_stock_engine.py — Kalkulasi Safety Stock & DDMRP',
    '  6.5 inventory_engine.py — Klasifikasi ABC-XYZ',
    '  6.6 rebalancing_engine.py — Optimasi Redistribusi Stok',
    '  6.7 landed_cost_engine.py — Kalkulasi Biaya Impor',
    '  6.8 control_tower_engine.py — Dashboard Kesehatan SCM',
    '  6.9 ddmrp_engine.py — DDMRP Lengkap',
    '  6.10 route_optimization_engine.py — Optimasi Rute Distribusi',
    '  6.11 wh_trans_mp_service.py — Simulasi Jaringan Hub',
    '  6.12 Utilitas (validators, imputation, llm_agent)',
    'Bab 7: Kesimpulan & Ringkasan',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('  ') else 'List Bullet')

# ═══════════════════════════════════════════
# BAB 1: GAMBARAN BESAR
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 1: Gambaran Besar Sistem', level=1)
doc.add_heading('1.1 Apa yang Dilakukan Sistem Ini?', level=2)
doc.add_paragraph(
    'Bayangkan Anda mengelola 28 gudang distribusi yang tersebar di seluruh Indonesia, '
    'dari Medan hingga Jayapura. Setiap gudang menerima barang dari pabrik, menyimpannya, '
    'dan mengirimkannya ke ribuan toko. Sistem WMS (Warehouse Management System) ini adalah '
    '"otak" digital yang membantu Anda menjawab pertanyaan-pertanyaan kritis berikut:'
)
questions = [
    ('Kapan harus pesan barang lagi?', 'Modul Forecast Engine dan Safety Stock Engine menghitung titik pesan ulang (Reorder Point) berdasarkan pola penjualan historis.'),
    ('Apakah gudang saya akan penuh?', 'Modul Occupancy Engine menghitung proyeksi kapasitas gudang minggu per minggu, dan memberi peringatan jika mendekati 100%.'),
    ('Barang mana yang paling penting?', 'Modul Inventory Engine mengklasifikasikan barang ke dalam kelompok ABC (berdasarkan nilai) dan XYZ (berdasarkan variabilitas penjualan).'),
    ('Berapa biaya kirim paling murah?', 'Modul Route Optimization Engine menggunakan algoritma seperti Clarke-Wright dan Genetic Algorithm untuk mencari rute pengiriman termurah.'),
    ('Gudang mana kelebihan/kekurangan stok?', 'Modul Rebalancing Engine merekomendasikan transfer stok antar gudang untuk menyeimbangkan persediaan.'),
    ('Di mana harus buka gudang baru?', 'Modul WH-TRANS-MP menggunakan K-Means Clustering untuk merekomendasikan lokasi hub distribusi optimal.'),
]
for q, a in questions:
    doc.add_paragraph(f'{q}', style='List Bullet')
    doc.add_paragraph(f'→ {a}')

doc.add_heading('1.2 Alur Kerja (Input → Proses → Output)', level=2)
doc.add_paragraph(
    'INPUT: Pengguna mengunggah file Excel/CSV berisi data penjualan, stok, kapasitas gudang, '
    'dan parameter biaya melalui antarmuka web (frontend Next.js).\n\n'
    'PROSES: Data dikirim ke server Python (FastAPI) yang menjalankan algoritma statistik, '
    'machine learning, dan optimasi. Server memproses data dan menghasilkan insight.\n\n'
    'OUTPUT: Hasil dikembalikan ke frontend berupa grafik interaktif, tabel, peta, KPI dashboard, '
    'dan file Excel yang bisa diunduh.'
)

# ═══════════════════════════════════════════
# BAB 2: ARSITEKTUR
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 2: Arsitektur & Struktur Kode', level=1)
doc.add_heading('2.1 Pola Arsitektur', level=2)
doc.add_paragraph(
    'Sistem ini mengikuti arsitektur Layered Monolith dengan pemisahan tanggung jawab:\n\n'
    '• Lapisan 1 (Routers): Menerima permintaan HTTP dari frontend. '
    'Ibarat resepsionis yang menerima telepon dan mengarahkan ke departemen yang tepat.\n\n'
    '• Lapisan 2 (Services/Engines): Berisi logika bisnis utama. '
    'Ibarat para ahli (analis, insinyur) yang melakukan pekerjaan analisis sesungguhnya.\n\n'
    '• Lapisan 3 (Utils): Fungsi-fungsi pembantu seperti validasi data dan pembersihan. '
    'Ibarat quality control yang memastikan data masuk sudah bersih dan benar.\n\n'
    '• Lapisan 4 (Schemas): Definisi struktur data (saat ini masih kosong, bisa dikembangkan).'
)

doc.add_heading('2.2 Peta File', level=2)
files_map = [
    ('main.py', 'Titik masuk aplikasi. Mendaftarkan semua modul ke server FastAPI.'),
    ('routers/*.py', '13 file router — menerima request HTTP dan meneruskan ke service.'),
    ('services/*.py', '12 file engine — berisi seluruh logika bisnis dan kalkulasi.'),
    ('utils/validators.py', 'Memvalidasi struktur kolom file Excel yang diunggah.'),
    ('utils/imputation.py', 'Membersihkan data kotor (null, format salah, dll).'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'File/Folder'
hdr[1].text = 'Fungsi'
for f, desc in files_map:
    row = table.add_row().cells
    row[0].text = f
    row[1].text = desc

# ═══════════════════════════════════════════
# BAB 3: CODE REVIEW
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 3: Code Review — Temuan & Rekomendasi', level=1)

doc.add_heading('3.1 Temuan Positif ✅', level=2)
positives = [
    'Pemisahan layer (Router vs Service) sudah konsisten — setiap router hanya memanggil service, tidak mengandung logika bisnis langsung.',
    'Fungsi _safe_float() digunakan secara konsisten di seluruh modul untuk mencegah error NaN/Inf yang bisa merusak respons JSON.',
    'Validasi kolom file upload dilakukan sebelum pemrosesan (validators.py), mencegah crash akibat kolom yang hilang.',
    'Algoritma forecasting menyertakan 16 model dan melakukan backtesting otomatis untuk memilih model terbaik per grup.',
    'DDMRP Engine mengimplementasikan metodologi lengkap (ADU, CoV, Buffer Zones, Net Flow) sesuai literatur Ptak & Smith.',
    'Imputation data bulan Indonesia (Januari→January, Mei→May, dst) sangat kontekstual dan berguna.',
    'Route Optimization Engine mengimplementasikan 3 algoritma (Nearest Neighbor, Clarke-Wright, Genetic Algorithm) dengan benchmarking.',
]
for p in positives:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('3.2 Temuan yang Perlu Diperbaiki ⚠️', level=2)
issues = [
    ('REDUNDANSI: Fungsi _safe_float() diduplikasi di 8 file',
     'Pindahkan ke satu lokasi di utils/helpers.py dan import dari sana. Ini mengurangi maintenance burden dan memastikan perilaku konsisten.',
     'MEDIUM'),
    ('REDUNDANSI: Mapping bulan Indonesia di imputation.py DAN forecast_engine.py',
     'Data duplikat rawan inkonsistensi. Buat satu CONST di utils/constants.py.',
     'LOW'),
    ('CORS allow_origins=["*"] di main.py',
     'Untuk produksi, batasi ke domain frontend aktual (misalnya http://localhost:3000). Wildcard * membuka risiko CSRF.',
     'HIGH'),
    ('Import di dalam fungsi (line 558, 658 di occupancy_engine.py)',
     'Circular import dihindari dengan lazy import. Lebih baik refactor agar services saling independen, atau gunakan dependency injection.',
     'MEDIUM'),
    ('GZipMiddleware import setelah app creation (main.py line 9)',
     'Pindahkan semua import ke bagian atas file sesuai PEP 8.',
     'LOW'),
    ('Schemas folder kosong (__init__.py saja)',
     'Manfaatkan Pydantic BaseModel untuk mendefinisikan request/response schema secara eksplisit, meningkatkan dokumentasi API otomatis (Swagger).',
     'MEDIUM'),
    ('Tidak ada logging terstruktur di service layer',
     'Tambahkan Python logging (bukan print) di setiap service engine untuk tracing error di produksi.',
     'HIGH'),
    ('Tidak ada rate limiting atau autentikasi',
     'Tambahkan middleware autentikasi (JWT/API Key) dan rate limiter sebelum deployment produksi.',
     'HIGH'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Temuan'
hdr[1].text = 'Rekomendasi'
hdr[2].text = 'Prioritas'
for issue, rec, prio in issues:
    row = table.add_row().cells
    row[0].text = issue
    row[1].text = rec
    row[2].text = prio

# ═══════════════════════════════════════════
# BAB 4: BUG HUNTING
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 4: Bug Hunting & Edge Cases', level=1)

doc.add_heading('4.1 Bug Aktual yang Ditemukan 🐛', level=2)
bugs = [
    ('forecast_engine.py line 40 — _compute_exog_factor DUPLIKAT',
     'Fungsi ini didefinisikan dua kali (line 40-56 dan line 605-623). Python akan menggunakan versi terakhir, tetapi ini membingungkan dan rawan konflik.',
     'Hapus salah satu definisi (lebih baik hapus yang pertama karena versi kedua memiliki pengecekan NaN tambahan).'),
    ('occupancy_engine.py — Penggunaan id(rec) sebagai dictionary key',
     'id() mengembalikan alamat memori objek. Jika objek di-garbage-collect dan alamatnya digunakan ulang, key akan tabrakan. Pada runtime normal ini aman karena semua records hidup bersamaan, tetapi ini adalah antipattern.',
     'Gunakan index integer atau tuple (cabang, grup, category) sebagai key.'),
    ('forecast_engine.py — ARIMAX proxy (line 242-243)',
     'Faktor exogenous hardcoded 1.02 (exogenous_factor = 1.02 * exog_factors[i]) berarti forecast selalu di-inflate 2%, yang merupakan bias sistematik.',
     'Hapus hardcoded 1.02 atau buat configurable.'),
    ('wh_trans_mp_service.py — parse_wh_trans_file import io dan pandas kedua kali',
     'Modul ini memiliki import duplikat di bagian bawah file (dari operasi append). Tidak error, tapi tidak bersih.',
     'Hapus import duplikat di bagian bawah file.'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Bug/Issue'
hdr[1].text = 'Dampak'
hdr[2].text = 'Solusi'
for bug, impact, fix in bugs:
    row = table.add_row().cells
    row[0].text = bug
    row[1].text = impact
    row[2].text = fix

doc.add_heading('4.2 Edge Cases yang Perlu Diuji', level=2)
edge_cases = [
    'File Excel kosong (0 baris data, hanya header) — apakah sistem mengembalikan error yang user-friendly?',
    'File Excel dengan kolom bernama berbeda (misalnya "cabang" vs "Cabang" vs "CABANG") — normalisasi sudah dilakukan di beberapa engine, tapi belum semua.',
    'Semua nilai penjualan = 0 — MAPE akan menghasilkan division by zero (sudah ditangani dengan mask, tapi perlu diverifikasi di semua 16 model).',
    'Data sangat sedikit (< 3 baris) — beberapa model forecast membutuhkan minimal data tertentu.',
    'Kapasitas gudang = 0 atau negatif — occupancy calculation akan menghasilkan None/Inf.',
    'Lead Time = 0 — sqrt(0) = 0, sehingga safety stock = 0 (ini bisa benar, tetapi perlu dikonfirmasi konteks bisnisnya).',
    'File upload > 100 MB — perlu batas ukuran file di FastAPI.',
    'Red Zone radius = 0 pada WH-TRANS-MP — setiap hub akan selalu "inside" the zone.',
    'Concurrent requests — FastAPI async, tetapi engine menggunakan pandas/numpy yang CPU-bound. Perlu background worker (Celery) untuk tugas berat.',
]
for ec in edge_cases:
    doc.add_paragraph(ec, style='List Bullet')

doc.add_heading('4.3 Rekomendasi Skenario Unit Test', level=2)
tests = [
    ('test_safe_float', 'Uji _safe_float dengan input: None, "abc", float("inf"), float("nan"), -0.0, 999999999999'),
    ('test_mape_zero_actual', 'Uji _mape ketika semua actual = 0 (harus return 0, bukan error)'),
    ('test_forecast_minimum_data', 'Uji setiap model forecast dengan data minimal (1, 2, 3 baris)'),
    ('test_occupancy_zero_capacity', 'Uji compute_occupancy ketika kapasitas gudang = 0'),
    ('test_haversine_same_point', 'Uji haversine(lat, lon, lat, lon) — harus return 0.0'),
    ('test_anti_gravity_on_center', 'Uji apply_anti_gravity ketika hub tepat di pusat Red Zone (dist=0)'),
    ('test_excel_missing_sheet', 'Uji parse file Excel tanpa sheet "Demand" — harus raise ValueError yang jelas'),
    ('test_abc_all_equal_volume', 'Uji klasifikasi ABC ketika semua item memiliki volume identik'),
    ('test_rebalancing_no_freight', 'Uji rebalancing ketika freight_df kosong — harus mengembalikan pesan error yang user-friendly'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Nama Test'
hdr[1].text = 'Skenario'
for name, scenario in tests:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = scenario

# ═══════════════════════════════════════════
# BAB 5: ANALISIS AKURASI PERHITUNGAN
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 5: Analisis Akurasi Perhitungan', level=1)

doc.add_heading('5.1 MAPE (Mean Absolute Percentage Error)', level=2)
doc.add_paragraph(
    'Rumus: MAPE = (1/n) × Σ|actual - predicted| / |actual| × 100%\n\n'
    'IMPLEMENTASI (forecast_engine.py line 10-15):\n'
    'Sudah BENAR. Fungsi _mape() menggunakan mask untuk mengabaikan data aktual yang bernilai 0, '
    'sehingga tidak terjadi division by zero. Ini adalah best practice standar.\n\n'
    'CATATAN: MAPE memiliki kelemahan bawaan — ia asimetris (overprediction dan underprediction diperlakukan berbeda). '
    'Untuk industri logistik, pertimbangkan menambahkan sMAPE (Symmetric MAPE) atau MASE sebagai metrik pelengkap.'
)

doc.add_heading('5.2 Safety Stock & Reorder Point', level=2)
doc.add_paragraph(
    'Rumus: SS = Z × σ × √LT\n'
    'Rumus: ROP = (ADU × LT) + SS\n\n'
    'IMPLEMENTASI (safety_stock_engine.py line 148-153):\n'
    'Sudah BENAR dan sesuai teori standar. Z-score diambil dari tabel lookup yang sudah diisi dengan benar '
    '(95% → 1.645, 99% → 2.326). Namun:\n\n'
    '• σ (std_usage) dihitung dengan ddof=1 (default pandas), yang berarti menggunakan sample standard deviation. '
    'Ini sudah tepat untuk estimasi populasi dari sample.\n\n'
    '• Lead Time diasumsikan konstan (deterministik). Jika lead time bervariasi, rumus seharusnya:\n'
    'SS = Z × √(LT × σ_demand² + ADU² × σ_leadtime²)\n'
    'Ini disebut formula safety stock dengan lead time uncertainty.'
)

doc.add_heading('5.3 Haversine Distance', level=2)
doc.add_paragraph(
    'Rumus: d = R × 2 × atan2(√a, √(1-a))\n'
    'dimana a = sin²(Δφ/2) + cos(φ₁) × cos(φ₂) × sin²(Δλ/2)\n\n'
    'IMPLEMENTASI (route_optimization_engine.py line 30-48 DAN wh_trans_mp_service.py line 9-22):\n'
    'Kedua implementasi BENAR secara matematis dan menghasilkan nilai identik. '
    'R = 6371.0 km (rata-rata jari-jari Bumi) sudah sesuai standar.\n\n'
    'CATATAN: Haversine mengasumsikan Bumi adalah bola sempurna. Untuk akurasi <0.5% di Indonesia, '
    'ini sudah cukup. Untuk presisi lebih tinggi, gunakan Vincenty formula.'
)

doc.add_heading('5.4 DDMRP Buffer Zones', level=2)
doc.add_paragraph(
    'IMPLEMENTASI (ddmrp_engine.py):\n'
    '• Red Zone = ADU × DLT × Variability_Factor — BENAR sesuai Ptak & Smith.\n'
    '• Yellow Zone = ADU × DLT — BENAR (coverage selama lead time).\n'
    '• Green Zone = max(ADU × Order_Cycle, MOQ) — BENAR.\n'
    '• Net Flow = On-Hand + On-Order − Qualified_Demand — BENAR.\n\n'
    'CATATAN PRESISI FLOATING-POINT:\n'
    'Semua kalkulasi menggunakan Python float (IEEE 754 double precision, 64-bit). '
    'Ini memberikan presisi sekitar 15-16 digit signifikan, yang lebih dari cukup untuk '
    'volume logistik skala Indonesia (maksimal ~10^12 unit/rupiah). Tidak ada risiko overflow atau '
    'kehilangan presisi pada skala operasi normal.'
)

doc.add_heading('5.5 Occupancy Calculation', level=2)
doc.add_paragraph(
    'Rumus: Occupancy(%) = Total_Stok_Positif / Kapasitas_Gudang × 100%\n\n'
    'IMPLEMENTASI (occupancy_engine.py line 198-212):\n'
    'Sudah BENAR. Menggunakan max(0.0, balance) sehingga stok negatif (shortage) '
    'tidak membebaskan ruang fisik gudang — ini logika bisnis yang tepat karena '
    'shortage pada satu kategori tidak mengurangi pemakaian ruang oleh kategori lain.\n\n'
    'REKOMENDASI: Pertimbangkan menambahkan volumetric factor (CBM per unit) jika '
    'barang-barang memiliki ukuran fisik yang sangat berbeda.'
)

# ═══════════════════════════════════════════
# BAB 6: BEDAH KODE PER MODUL
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 6: Bedah Kode Per Modul (Line-by-Line)', level=1)

# 6.1 main.py
doc.add_heading('6.1 main.py — Pintu Masuk Aplikasi', level=2)
doc.add_paragraph(
    'Analoginya: File ini seperti pintu utama sebuah gedung perkantoran. '
    'Ia menerima semua tamu (request dari browser), lalu mengarahkan mereka ke lantai/departemen yang tepat.'
)
add_code_table(doc, [
    ('from fastapi import FastAPI', 'Mengimpor "kerangka bangunan" web server bernama FastAPI. Ibarat membeli pondasi gedung sebelum membangun.'),
    ('from routers import occupancy, forecast, ...', 'Mengundang semua "departemen" (modul) untuk bergabung ke gedung. Setiap router adalah satu departemen spesialis.'),
    ('app = FastAPI(title="Demand Supply Planning API")', 'Mendirikan gedung dan memberinya nama "Demand Supply Planning API". Variabel app adalah gedung itu sendiri.'),
    ('app.add_middleware(CORSMiddleware, allow_origins=["*"])', 'Memasang "satpam" yang mengizinkan siapa saja masuk. Tanda "*" berarti semua orang dibolehkan — ini perlu dibatasi untuk produksi.'),
    ('app.add_middleware(GZipMiddleware, minimum_size=1000)', 'Memasang "kompresor udara" — data yang dikirim ke browser akan dimampatkan agar lebih cepat jika ukurannya > 1000 byte.'),
    ('app.include_router(occupancy.router, prefix="/api/v1")', 'Mendaftarkan departemen Occupancy di alamat /api/v1/... — semua request ke alamat ini akan diteruskan ke modul occupancy.'),
    ('@app.get("/")', 'Membuat "papan nama" di pintu utama. Jika seseorang mengakses alamat dasar, mereka melihat pesan "API is running".'),
    ('@app.get("/health")', 'Membuat "kotak P3K" — endpoint untuk memeriksa apakah server masih hidup (health check). Biasa digunakan oleh monitoring tools.'),
])

# 6.2 occupancy_engine.py
doc.add_heading('6.2 occupancy_engine.py — Mesin MRP & Kapasitas Gudang', level=2)
doc.add_paragraph(
    'Modul ini adalah jantung dari kalkulasi Material Requirements Planning (MRP). '
    'Analogi: Bayangkan seorang manajer gudang yang setiap minggu menghitung:\n'
    '"Berapa stok saya sekarang? Berapa yang akan datang? Berapa yang akan keluar? '
    'Apakah gudang saya cukup besar?"\n\n'
    'Modul ini melakukan semua kalkulasi itu secara otomatis dari file Excel.'
)
add_code_table(doc, [
    ('def _safe_float(v):', 'Fungsi "pengaman". Jika ada data yang bukan angka (misalnya teks, kosong, atau tak terhingga), fungsi ini mengubahnya menjadi 0.0 agar tidak crash.'),
    ('FIXED_RAW_COLS = 4', 'Konstanta: 4 kolom pertama di Excel adalah data tetap (No, Cabang, Grup, Category). Ibarat 4 kolom paling kiri yang tidak berubah.'),
    ('COLS_PER_WEEK = 4', 'Setiap minggu memiliki 4 sub-kolom data: TO (Transfer Order), Vessel (Pengiriman Kapal), Forecast (Prediksi), Target.'),
    ('class RawRecord:', 'Blueprint/cetakan untuk menyimpan data 1 baris Excel. Ibarat satu "kartu identitas" barang yang menyimpan semua informasinya.'),
    ('def compute_balance_series(record, n_weeks, demand_kind):', 'FUNGSI INTI. Menghitung saldo stok minggu per minggu.\n'
     'Minggu 1: Saldo = Stok_Awal + TO + Vessel - Demand\n'
     'Minggu 2+: Saldo = Saldo_Minggu_Sebelumnya + TO + Vessel - Demand\n'
     'Ini seperti buku tabungan: setiap minggu ada pemasukan (TO, Vessel) dan pengeluaran (Demand).'),
    ('def compute_ratio_series(balances, record, ...):', 'Menghitung rasio kecukupan stok. Jika rasio < 1 (100%), artinya stok tidak cukup untuk memenuhi permintaan minggu depan.'),
    ('def compute_occupancy(records, balances, wh_capacity):', 'Menghitung persentase ruang gudang yang terpakai. Occupancy = Total_Stok / Kapasitas × 100%. Jika > 100%, gudang terlalu penuh.'),
    ('def generate_insights(records, bal_f, bal_t, ...):', 'Menghasilkan rekomendasi otomatis. Misalnya: "RISIKO KEKURANGAN — Cabang Jakarta periode MAR-2, balance hanya 45% dari kebutuhan".'),
])

# 6.3 forecast_engine.py
doc.add_heading('6.3 forecast_engine.py — Mesin Peramalan 16 Model', level=2)
doc.add_paragraph(
    'Modul ini seperti tim riset yang berisi 16 orang analis. Masing-masing menggunakan metode berbeda '
    'untuk memprediksi penjualan masa depan, lalu satu "juri" memilih analis mana yang paling akurat.'
)
add_code_table(doc, [
    ('def _mape(y_true, y_pred):', 'Menghitung seberapa akurat prediksi dalam persen. MAPE 10% berarti rata-rata prediksi meleset 10% dari kenyataan. Semakin kecil, semakin bagus.'),
    ('def _ses_forecast(y_train, steps, alpha=0.3):', 'Model SES (Simple Exponential Smoothing). Ibarat meramal cuaca hanya berdasarkan cuaca hari ini — jika kemarin panas, besok mungkin juga panas. Alpha=0.3 berarti 30% bobot di data terbaru, 70% di data lama.'),
    ('def _hw_forecast(y_train, steps, seasonality=3):', 'Model Holt-Winters. Lebih canggih dari SES karena mengenali "musim" (seasonal pattern). Misalnya penjualan selalu naik tiap 3 bulan.'),
    ('def _gb_forecast(y_train, steps, exog_train=None):', 'Model Gradient Boosting (proxy XGBoost). Membuat 5 "pohon keputusan kecil" yang secara bertahap memperbaiki error prediksi. Setiap pohon belajar dari kesalahan pohon sebelumnya.'),
    ('def _bilstm_proxy_forecast(y_train, steps):', 'Model BiLSTM (Bidirectional LSTM). Mensimulasikan jaringan saraf tiruan yang membaca data maju DAN mundur. Ibarat membaca buku dari depan dan belakang sekaligus.'),
    ('def _arimax_proxy_forecast(y_train, steps, exog_train):', 'Model ARIMAX. Menggabungkan autoregression (data masa lalu), moving average (rata-rata error), dan faktor eksternal (AO, RO, NOO). PERINGATAN: ada bias +2% hardcoded.'),
    ('best_model_info = min(models_eval, key=lambda x: (x["mape"], x["rmse"]))', 'JURI! Memilih model dengan MAPE terendah (dan RMSE sebagai tiebreaker). Model ini yang akan digunakan untuk prediksi masa depan.'),
    ('safety_stock = 1.65 * std_dev', 'Menghitung Safety Stock = Z × σ. Z=1.65 setara service level 95% (artinya 95% kemungkinan stok cukup). σ = standar deviasi penjualan historis.'),
    ('rop = avg_sales + safety_stock', 'Reorder Point = rata-rata penjualan + safety stock. Jika stok turun di bawah angka ini, SEGERA pesan ulang.'),
])

# 6.4-6.5
doc.add_heading('6.4 safety_stock_engine.py — Kalkulasi Safety Stock & DDMRP', level=2)
doc.add_paragraph('Modul ini menghitung buffer stok pengaman (Safety Stock) dan zona DDMRP untuk setiap SKU di setiap cabang.')
add_code_table(doc, [
    ('Z_SCORE_TABLE = {0.90: 1.282, ..., 0.99: 2.326}', 'Tabel konversi Service Level → Z-Score. 95% berarti "kita ingin 95% yakin stok cukup". Z-Score 1.645 menunjukkan berapa standar deviasi kita perlu buffer.'),
    ('safety_stock = z_score * std_usage * math.sqrt(lt)', 'RUMUS INTI: SS = Z × σ × √LT. Semakin bervariasi penjualan (σ besar) DAN semakin lama waktu tunggu pengiriman (LT), semakin besar stok pengaman yang diperlukan.'),
    ('rop = (adu * lt) + safety_stock', 'Reorder Point = kebutuhan selama lead time + safety stock. Jika stok ≤ ROP, saatnya pesan ulang!'),
    ('red_zone = yellow_zone * lt_variability_factor', 'Zona Merah DDMRP: buffer darurat. Semakin volatile demand, semakin tebal zona merah.'),
    ('green_zone = max(adu * order_cycle, moq)', 'Zona Hijau DDMRP: batas pemesanan. Minimal harus sebesar MOQ (Minimum Order Quantity) agar pesanan terpenuhi.'),
    ('net_flow = current_stock + in_transit - backorder', 'Posisi Net Flow = stok di tangan + yang sedang dikirim - pesanan yang belum terpenuhi. Ini menentukan apakah perlu pesan atau tidak.'),
])

doc.add_heading('6.5 inventory_engine.py — Klasifikasi ABC-XYZ', level=2)
doc.add_paragraph('Modul ini mengklasifikasikan produk berdasarkan dua dimensi: nilai kontribusi (ABC) dan variabilitas permintaan (XYZ).')
add_code_table(doc, [
    ('cv = std_sales / mean_sales', 'Coefficient of Variation (CoV). Mengukur seberapa "liar" penjualan suatu barang. Jika CV < 0.5 → stabil (X), 0.5-1.0 → moderat (Y), > 1.0 → sangat liar (Z).'),
    ('get_abc(perc): if perc <= 0.80 → A', 'Klasifikasi ABC berdasarkan kontribusi kumulatif. Barang yang menyumbang 0-80% total penjualan = Kelas A (paling penting), 80-95% = B, sisanya = C.'),
    ('doh = max(0.0, current_on_hand) / daily_sales', 'Days on Hand: berapa hari stok akan bertahan. Jika > 90 hari, barang dikategorikan sebagai Dead Stock (tersimpan terlalu lama).'),
    ('_recommend_strategy(abc, xyz, ...)', 'Matriks rekomendasi 3×3. Contoh: AX = "tight safety stock" (barang mahal, penjualan stabil), CZ = "candidate for discontinuation" (barang murah, penjualan tidak terprediksi).'),
])

# 6.6-6.12 (ringkas agar dokumen tidak terlalu panjang)
for section_num, (title, desc, key_funcs) in enumerate([
    ('6.6 rebalancing_engine.py — Redistribusi Stok Antar Gudang',
     'Mengoptimalkan transfer stok dari cabang kelebihan ke cabang kekurangan, mempertimbangkan biaya freight.',
     [('analyze_rebalancing(stock_df, demand_df, freight_df)', 'Fungsi utama yang mencocokan surplus stok di satu cabang dengan defisit di cabang lain, meminimalkan total biaya pengiriman.')]),
    ('6.7 landed_cost_engine.py — Kalkulasi Biaya Impor Kontainer',
     'Menghitung total biaya yang dikeluarkan untuk mendatangkan barang impor, dari FOB hingga sampai di gudang (landed cost).',
     [('analyze_landed_cost(tracking_df, allocation_df, exchange_rate)', 'Menerima data tracking kontainer (BL, ETA, biaya freight/duty/THC) dan alokasi SKU, lalu menghitung total landed cost per unit dalam Rupiah.')]),
    ('6.8 control_tower_engine.py — Dashboard Kesehatan SCM 28 Cabang',
     'Memberikan dashboard "bird-eye view" kesehatan supply chain seluruh cabang, termasuk In-Stock Rate, OTIF Score, dan Days of Supply.',
     [('analyze_control_tower(df)', 'Menganalisis data performa setiap cabang, menentukan skor kesehatan, dan menghasilkan alert untuk cabang yang berisiko.')]),
    ('6.9 ddmrp_engine.py — DDMRP Lengkap (664 baris)',
     'Implementasi penuh metodologi Demand Driven MRP, termasuk ADU, CoV, Buffer Zones, Net Flow, dan simulasi Monte Carlo.',
     [('calc_adu(sales, period_days=30)', 'Average Daily Usage = Total Penjualan / Total Hari. Ini adalah heartbeat dari DDMRP.'),
      ('classify_variability_factor(cov)', 'Mengubah CoV menjadi kategori (Low/Medium/High) dan faktor pengali buffer yang sesuai.')]),
    ('6.10 route_optimization_engine.py — Optimasi Rute (1139 baris)',
     'Modul terbesar yang mengimplementasikan 3 algoritma optimasi rute: Nearest Neighbor, Clarke-Wright Savings, dan Genetic Algorithm.',
     [('haversine_km(lat1, lon1, lat2, lon2)', 'Menghitung jarak antar titik GPS menggunakan rumus Haversine (bola Bumi).'),
      ('clarke_wright_savings(locations, dist_matrix, ...)', 'Algoritma penghematan klasik — menggabungkan rute individual menjadi rute bersama jika menghemat jarak.'),
      ('genetic_algorithm(locations, dist_matrix, ...)', 'Simulasi evolusi biologi — rute-rute terbaik "berkembang biak" dan bermutasi untuk menghasilkan solusi lebih baik.')]),
    ('6.11 wh_trans_mp_service.py — Simulasi Hub Desentralisasi',
     'Menggunakan K-Means Clustering untuk menemukan lokasi hub optimal, dengan logika "anti-gravitasi" untuk menghindari zona mahal/macet.',
     [('haversine(lat1, lon1, lat2, lon2)', 'Menghitung jarak antara dua titik GPS dalam kilometer.'),
      ('apply_anti_gravity(hub_lat, hub_lon, red_zones)', 'Jika hub jatuh di dalam Red Zone, fungsi ini "mendorongnya keluar" ke jarak aman + 1km buffer.'),
      ('simulate_network(data, num_hubs=3)', 'Menjalankan K-Means, menerapkan anti-gravity, lalu menghitung penghematan biaya jika menggunakan N hub vs 1 hub sentral.')]),
    ('6.12 Utilitas — validators.py, imputation.py, llm_agent.py',
     'Fungsi pendukung untuk validasi, pembersihan data, dan chatbot sederhana.',
     [('validate_occupancy_schema(df)', 'Memastikan file Excel memiliki kolom wajib (Cabang, Category, On Hand, In, Out, Capacity, Date).'),
      ('clean_forecast_data(df)', 'Membersihkan data: mengubah nama bulan Indonesia ke Inggris, mengkonversi teks ke tanggal, mengisi nilai kosong dengan 0.'),
      ('process_chat_query(query, context_data)', 'Chatbot berbasis aturan (bukan AI sesungguhnya). Mencocokkan kata kunci dalam pertanyaan pengguna untuk memberikan jawaban dari data yang sudah dianalisis.')]),
], start=1):
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    add_code_table(doc, key_funcs)

# ═══════════════════════════════════════════
# BAB 7: KESIMPULAN
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Bab 7: Kesimpulan & Ringkasan', level=1)

doc.add_heading('7.1 Ringkasan Arsitektur', level=2)
doc.add_paragraph(
    'Sistem WMS Monorepo ini terdiri dari 12 modul service Python yang saling terintegrasi, '
    'melayani 13 endpoint API melalui FastAPI. Total ~4,500+ baris kode Python murni (tidak termasuk venv). '
    'Arsitekturnya mengikuti pola Layered Monolith yang bersih dengan pemisahan Router-Service-Utils.'
)

doc.add_heading('7.2 Kekuatan Utama', level=2)
strengths = [
    'Kelengkapan fitur SCM: Dari forecasting hingga route optimization, mencakup hampir seluruh siklus supply chain.',
    'Implementasi 16 model forecasting dengan auto-selection berdasarkan backtesting — jarang ditemukan di sistem serupa.',
    'DDMRP compliance: Implementasi sesuai literatur akademik (Ptak & Smith).',
    'Error handling defensif: _safe_float() mencegah crash JSON serialization.',
    'Kontekstualisasi Indonesia: Mapping bulan Indonesia, 28 cabang, mata uang Rupiah.',
]
for s in strengths:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('7.3 Area Perbaikan Prioritas', level=2)
priorities = [
    'PRIORITAS 1 (Keamanan): Batasi CORS, tambahkan autentikasi API.',
    'PRIORITAS 2 (Kualitas Kode): Eliminasi duplikasi _safe_float() dan _compute_exog_factor().',
    'PRIORITAS 3 (Akurasi): Hapus bias +2% di ARIMAX proxy, tambahkan sMAPE sebagai metrik pelengkap.',
    'PRIORITAS 4 (Skalabilitas): Tambahkan background worker untuk tugas berat (forecasting 1000+ SKU).',
    'PRIORITAS 5 (Testing): Implementasikan unit test suite minimal (9 skenario yang direkomendasikan di Bab 4).',
]
for p in priorities:
    doc.add_paragraph(p, style='List Number')

doc.add_heading('7.4 Penutup', level=2)
doc.add_paragraph(
    'Sistem ini sudah memiliki fondasi yang sangat baik untuk sebuah platform analitik supply chain. '
    'Dengan perbaikan pada area keamanan, testing, dan eliminasi redundansi, sistem ini siap untuk '
    'deployment produksi yang melayani operasi logistik skala nasional.'
)

# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════
output_path = r'c:\Users\DELL\Downloads\python\Vibes Coding\wms-monorepo\frontend\public\WMS_Dokumentasi_Komprehensif.docx'
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
